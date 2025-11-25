"""
Export test cases to Word document (.docx)
Requires: python-docx
Install: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import json

def create_test_case_word(test_cases, output_file="test_cases.docx"):
    """
    Convert test cases to a formatted Word document
    
    Args:
        test_cases: List of dictionaries containing test case information
        output_file: Output filename for the Word document
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('Test Cases Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Total Test Cases: {len(test_cases)}')
    doc.add_paragraph('')
    
    # Process each test case
    for idx, test_case in enumerate(test_cases, 1):
        # Test case header
        heading = doc.add_heading(f'Test Case #{idx}: {test_case.get("name", "Unnamed Test")}', level=1)
        
        # Create table for test case details
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Add test case fields
        fields = [
            ('ID', test_case.get('id', f'TC_{idx:03d}')),
            ('Description', test_case.get('description', 'N/A')),
            ('Preconditions', test_case.get('preconditions', 'N/A')),
            ('Test Steps', test_case.get('steps', 'N/A')),
            ('Expected Result', test_case.get('expected_result', 'N/A')),
            ('Actual Result', test_case.get('actual_result', '')),
            ('Status', test_case.get('status', 'Not Executed')),
            ('Priority', test_case.get('priority', 'Medium')),
            ('Test Type', test_case.get('test_type', 'Functional')),
        ]
        
        for field_name, field_value in fields:
            row = table.add_row()
            row.cells[0].text = field_name
            row.cells[1].text = str(field_value)
            # Bold the field names
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Add spacing
        doc.add_paragraph('')
    
    # Save document
    doc.save(output_file)
    print(f"✓ Word document saved: {output_file}")
    return output_file


# Example usage with sample test cases
if __name__ == "__main__":
    # Sample test cases structure (replace with your Copilot-generated test cases)
    sample_test_cases = [
        {
            "id": "TC_001",
            "name": "User Login with Valid Credentials",
            "description": "Verify that user can login with valid username and password",
            "preconditions": "User account exists in the system",
            "steps": "1. Navigate to login page\n2. Enter valid username\n3. Enter valid password\n4. Click Login button",
            "expected_result": "User is successfully logged in and redirected to dashboard",
            "status": "Pass",
            "priority": "High",
            "test_type": "Functional"
        },
        {
            "id": "TC_002",
            "name": "User Login with Invalid Password",
            "description": "Verify error message when invalid password is entered",
            "preconditions": "User account exists in the system",
            "steps": "1. Navigate to login page\n2. Enter valid username\n3. Enter invalid password\n4. Click Login button",
            "expected_result": "Error message 'Invalid credentials' is displayed",
            "status": "Pass",
            "priority": "High",
            "test_type": "Negative Testing"
        }
    ]
    
    # Create Word document
    create_test_case_word(sample_test_cases, "test_cases_output.docx")
